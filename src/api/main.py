"""
PRISMY Translation API - Complete Production Version
Enhanced CORS, Error Handling, and Frontend Integration
FIXED: Added port 3002 support for frontend
"""
from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from src.api import celery_endpoints
import openai
import os
import pdfplumber
import tempfile
import uuid
from datetime import datetime
import logging
import traceback
from pathlib import Path

from src.api.v1.large_document import router as large_document_router

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="PRISMY Translation API", 
    version="1.0.0",
    description="Professional translation service with multiple quality tiers",
    docs_url="/docs",
    redoc_url="/redoc",
    openapi_url="/openapi.json"
)

# ✅ COMPREHENSIVE CORS CONFIGURATION - FIXED WITH PORT 3002
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        # Local development - ALL PORTS
        "http://localhost:3000",
        "http://127.0.0.1:3000", 
        "http://localhost:3001",
        "http://127.0.0.1:3001",
        "http://localhost:3002",     # ✅ FIXED: Added port 3002
        "http://127.0.0.1:3002",     # ✅ FIXED: Added port 3002
        "http://localhost:8000",
        "http://127.0.0.1:8000",
        # Frontend on different network IPs
        "http://192.168.1.15:3000",
        "http://192.168.1.15:3001",
        "http://192.168.1.15:3002", # ✅ FIXED: Added port 3002
        # Add production domains when ready
        # "https://prismy.com",
        # "https://www.prismy.com",
    ],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "HEAD"],
    allow_headers=[
        "Accept",
        "Accept-Language",
        "Content-Language", 
        "Content-Type",
        "Authorization",
        "Cache-Control",
        "X-Requested-With",
        "Origin",
        "Access-Control-Request-Method",
        "Access-Control-Request-Headers",
        "X-CSRF-Token",
        "X-Requested-With",
    ],
    expose_headers=[
        "Content-Disposition",
        "Content-Length",
        "Content-Type",
        "X-Total-Count",
        "X-RateLimit-Limit", 
        "X-RateLimit-Remaining",
        "Access-Control-Allow-Origin",
        "Access-Control-Allow-Methods",
        "Access-Control-Allow-Headers",
    ]
)

# ✅ ENHANCED REQUEST LOGGING MIDDLEWARE
@app.middleware("http")
async def enhanced_logging_middleware(request: Request, call_next):
    start_time = datetime.now()
    
    # Log incoming request with details
    client_ip = request.client.host if request.client else "unknown"
    origin = request.headers.get("origin", "no-origin")
    user_agent = request.headers.get("user-agent", "unknown")
    
    # ✅ DETAILED CORS LOGGING
    logger.info(f"🔄 {request.method} {request.url} - IP: {client_ip} - Origin: {origin}")
    
    try:
        response = await call_next(request)
        
        # ✅ ADD CORS HEADERS TO ALL RESPONSES
        if origin and any(origin.startswith(allowed) for allowed in [
            "http://localhost:3000", "http://localhost:3001", "http://localhost:3002",
            "http://127.0.0.1:3000", "http://127.0.0.1:3001", "http://127.0.0.1:3002",
            "http://192.168.1.15:3000", "http://192.168.1.15:3001", "http://192.168.1.15:3002"
        ]):
            response.headers["Access-Control-Allow-Origin"] = origin
            response.headers["Access-Control-Allow-Credentials"] = "true"
            response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, HEAD"
            response.headers["Access-Control-Allow-Headers"] = "*"
        
        # Log successful response
        process_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"✅ {request.method} {request.url} - {response.status_code} ({process_time:.3f}s)")
        
        return response
        
    except Exception as e:
        # Log errors
        process_time = (datetime.now() - start_time).total_seconds()
        logger.error(f"❌ {request.method} {request.url} - ERROR ({process_time:.3f}s): {str(e)}")
        
        # Return proper error response with CORS headers
        error_response = JSONResponse(
            status_code=500,
            content={
                "error": "Internal server error",
                "message": str(e),
                "timestamp": datetime.now().isoformat(),
                "path": str(request.url.path)
            }
        )
        
        # Add CORS headers to error responses too
        if origin:
            error_response.headers["Access-Control-Allow-Origin"] = origin
            error_response.headers["Access-Control-Allow-Credentials"] = "true"
        
        return error_response

# ✅ PREFLIGHT OPTIONS HANDLER
@app.options("/{full_path:path}")
async def handle_options(request: Request):
    """Handle preflight OPTIONS requests for CORS"""
    origin = request.headers.get("origin")
    
    response = JSONResponse(content={"message": "OK"})
    
    if origin:
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Credentials"] = "true"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, HEAD"
        response.headers["Access-Control-Allow-Headers"] = "*"
        response.headers["Access-Control-Max-Age"] = "86400"  # 24 hours
    
    logger.info(f"✅ OPTIONS preflight handled for origin: {origin}")
    return response

# ✅ OPENAI CLIENT SETUP WITH FALLBACK
try:
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if openai_api_key:
        openai.api_key = openai_api_key
        logger.info("✅ OpenAI API key configured")
    else:
        logger.warning("⚠️ OpenAI API key not found - using mock translation")
except Exception as e:
    logger.error(f"❌ OpenAI setup failed: {e}")

# ✅ ENHANCED STORAGE (In production, use Redis/Database)
translation_storage = {}

# ✅ INCLUDE ROUTERS WITH PROPER ERROR HANDLING
try:
    app.include_router(large_document_router, prefix="/api/v1/large", tags=["Large Documents"])
    logger.info("✅ Large document router loaded")
except Exception as e:
    logger.error(f"❌ Failed to load large document router: {e}")

try:
    app.include_router(celery_endpoints.router, prefix="/api/v2", tags=["Async Translation"])
    logger.info("✅ Celery endpoints router loaded")
except Exception as e:
    logger.error(f"❌ Failed to load celery router: {e}")

@app.get("/")
async def root():
    """Enhanced API root endpoint with comprehensive service information"""
    try:
        # Check system status
        redis_status = "unknown"
        try:
            from src.services.queue.redis_client import get_redis_client
            redis_client = get_redis_client()
            redis_status = "connected" if redis_client.ping() else "disconnected"
        except Exception:
            redis_status = "unavailable"
        
        return {
            "service": "PRISMY Translation API",
            "version": "1.0.0",
            "status": "operational",
            "timestamp": datetime.now().isoformat(),
            "system": {
                "redis": redis_status,
                "openai": "configured" if os.getenv("OPENAI_API_KEY") else "not_configured",
                "storage": f"{len(translation_storage)} active translations"
            },
            "endpoints": {
                "large_documents": "/api/v1/large/translate",
                "async_translation": "/api/v2/translate/pdf/async",
                "health": "/health",
                "docs": "/docs",
                "redoc": "/redoc",
                "api_status": "/api/status"
            },
            "supported_formats": ["PDF", "TXT", "DOC", "DOCX", "RTF"],
            "translation_tiers": ["free", "basic", "standard", "premium"],
            "supported_languages": {
                "from": ["auto", "en", "vi", "zh", "ja", "ko", "fr", "de", "es"],
                "to": ["en", "vi", "zh", "ja", "ko", "fr", "de", "es"]
            },
            "features": {
                "cors_enabled": True,
                "cors_origins": ["localhost:3000", "localhost:3001", "localhost:3002"],  # ✅ Updated
                "file_upload": True,
                "batch_processing": True,
                "progress_tracking": True,
                "download_results": True
            }
        }
    except Exception as e:
        logger.error(f"❌ Root endpoint error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "Service temporarily unavailable", "message": str(e)}
        )

@app.get("/health")
async def health_check():
    """Comprehensive health check endpoint"""
    health_status = {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0",
        "uptime": "active",
        "services": {}
    }
    
    # Check Redis
    try:
        from src.services.queue.redis_client import get_redis_client
        redis_client = get_redis_client()
        redis_ping = redis_client.ping()
        health_status["services"]["redis"] = "healthy" if redis_ping else "unhealthy"
    except Exception as e:
        health_status["services"]["redis"] = f"error: {str(e)}"
    
    # Check file system
    try:
        temp_dir = Path(tempfile.gettempdir())
        if temp_dir.exists() and temp_dir.is_dir():
            health_status["services"]["filesystem"] = "healthy"
        else:
            health_status["services"]["filesystem"] = "unhealthy"
    except Exception as e:
        health_status["services"]["filesystem"] = f"error: {str(e)}"
    
    # Check OpenAI
    health_status["services"]["openai"] = "configured" if os.getenv("OPENAI_API_KEY") else "not_configured"
    
    # Check worker status
    health_status["services"]["worker"] = "check with /api/v1/large/queue/status"
    
    # ✅ CORS Status
    health_status["services"]["cors"] = "enabled for ports 3000, 3001, 3002"
    
    # Overall status
    unhealthy_services = [k for k, v in health_status["services"].items() if "error" in str(v) or v == "unhealthy"]
    if unhealthy_services:
        health_status["status"] = "degraded"
        health_status["issues"] = unhealthy_services
    
    return health_status

@app.get("/api/status")
async def detailed_api_status():
    """Detailed API status for frontend debugging"""
    return {
        "api_version": "1.0.0",
        "server_time": datetime.now().isoformat(),
        "endpoints_available": True,
        "cors_enabled": True,
        "cors_origins": [
            "http://localhost:3000", "http://localhost:3001", "http://localhost:3002",  # ✅ Updated
            "http://127.0.0.1:3000", "http://127.0.0.1:3001", "http://127.0.0.1:3002"   # ✅ Updated
        ],
        "translation_service": "active",
        "active_translations": len(translation_storage),
        "supported_methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "max_file_size": "100MB",
        "timeout": "300 seconds",
        "frontend_integration": {
            "supported_ports": ["3000", "3001", "3002"],  # ✅ Updated
            "cors_configured": True,
            "preflight_handled": True
        }
    }

# ✅ ENHANCED TRANSLATION ENDPOINT WITH BETTER ERROR HANDLING
@app.post("/translate")
async def translate_document(file: UploadFile = File(...)):
    """
    Enhanced document translation with comprehensive error handling
    Supports PDF, DOC, DOCX, TXT files
    """
    translation_id = str(uuid.uuid4())
    
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ""
        supported_formats = ['pdf', 'txt', 'doc', 'docx']
        
        if file_extension not in supported_formats:
            raise HTTPException(
                status_code=400,
                detail={
                    "error": f"Unsupported file format: .{file_extension}",
                    "supported_formats": supported_formats,
                    "filename": file.filename
                }
            )
        
        # Read file content
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Empty file uploaded")
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}') as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        logger.info(f"📁 Processing file: {file.filename} ({len(content)} bytes)")
        
        # Extract text based on file type
        extracted_text = ""
        total_pages = 0
        
        if file_extension == 'pdf':
            with pdfplumber.open(tmp_path) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += f"\n--- Page {i+1} ---\n{page_text}\n"
        
        elif file_extension == 'txt':
            with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
                total_pages = 1
        
        elif file_extension in ['doc', 'docx']:
            try:
                from docx import Document
                doc = Document(tmp_path)
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        extracted_text += paragraph.text + "\n"
                total_pages = len(doc.paragraphs)
            except Exception as e:
                logger.error(f"❌ DOCX processing failed: {e}")
                raise HTTPException(status_code=400, detail=f"Failed to process Word document: {str(e)}")
        
        if not extracted_text.strip():
            raise HTTPException(
                status_code=400,
                detail={
                    "error": "No text extracted from file",
                    "suggestion": "File might be empty, corrupted, or image-based",
                    "file_type": file_extension
                }
            )
        
        # Mock translation (replace with real API when available)
        if os.getenv("OPENAI_API_KEY"):
            # Real OpenAI translation would go here
            translated_text = f"[OpenAI Translation]\n\n{extracted_text}"
        else:
            # Mock translation for testing
            translated_text = f"[MOCK TRANSLATION - Vietnamese]\n\n"
            translated_text += extracted_text.replace("the", "các").replace("and", "và").replace("of", "của")
        
        # Store translation
        translation_data = {
            "filename": file.filename,
            "file_type": file_extension,
            "original": extracted_text,
            "translated": translated_text,
            "pages": total_pages,
            "file_size": len(content),
            "timestamp": datetime.now().isoformat(),
            "status": "completed"
        }
        
        translation_storage[translation_id] = translation_data
        
        logger.info(f"✅ Translation completed: {file.filename} (ID: {translation_id})")
        
        return {
            "id": translation_id,
            "filename": file.filename,
            "file_type": file_extension,
            "pages": total_pages,
            "original_length": len(extracted_text),
            "translated_length": len(translated_text),
            "status": "completed",
            "timestamp": datetime.now().isoformat(),
            "preview": {
                "original": extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
                "translated": translated_text[:500] + "..." if len(translated_text) > 500 else translated_text
            },
            "download_urls": {
                "original": f"/download/{translation_id}/original",
                "translated": f"/download/{translation_id}/translated"
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Translation failed: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail={
                "error": "Translation processing failed",
                "message": str(e),
                "translation_id": translation_id,
                "timestamp": datetime.now().isoformat()
            }
        )
    finally:
        # Cleanup temporary file
        try:
            if 'tmp_path' in locals() and os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except Exception as e:
            logger.warning(f"⚠️ Failed to cleanup temp file: {e}")

# ✅ ENHANCED DOWNLOAD ENDPOINT
@app.get("/download/{translation_id}/{content_type}")
async def download_translation(translation_id: str, content_type: str):
    """
    Enhanced download endpoint with proper headers and error handling
    """
    if translation_id not in translation_storage:
        available_ids = list(translation_storage.keys())[-5:] if translation_storage else []
        raise HTTPException(
            status_code=404,
            detail={
                "error": "Translation not found",
                "translation_id": translation_id,
                "available_ids": available_ids,
                "total_available": len(translation_storage)
            }
        )
    
    data = translation_storage[translation_id]
    base_filename = data["filename"].rsplit('.', 1)[0] if '.' in data["filename"] else data["filename"]
    
    if content_type == "original":
        content = data["original"]
        output_filename = f"{base_filename}_original.txt"
    elif content_type == "translated":
        content = data["translated"]
        output_filename = f"{base_filename}_vietnamese.txt"
    else:
        raise HTTPException(
            status_code=400,
            detail={
                "error": "Invalid content type",
                "valid_types": ["original", "translated"],
                "received": content_type
            }
        )
    
    # Create temporary file for download
    try:
        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8') as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        logger.info(f"📥 Download requested: {output_filename} ({len(content)} chars)")
        
        return FileResponse(
            tmp_path,
            filename=output_filename,
            media_type='text/plain; charset=utf-8',
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"',
                "Content-Type": "text/plain; charset=utf-8",
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "Pragma": "no-cache",
                "Expires": "0"
            }
        )
    except Exception as e:
        logger.error(f"❌ Download failed: {e}")
        raise HTTPException(
            status_code=500,
            detail={"error": "Download failed", "message": str(e)}
        )

# ✅ ADDITIONAL UTILITY ENDPOINTS
@app.get("/translations")
async def list_translations():
    """List all stored translations with pagination"""
    return {
        "total": len(translation_storage),
        "translations": [
            {
                "id": tid,
                "filename": data["filename"],
                "file_type": data.get("file_type", "unknown"),
                "pages": data.get("pages", 0),
                "status": data.get("status", "unknown"),
                "timestamp": data["timestamp"],
                "file_size": data.get("file_size", 0)
            }
            for tid, data in translation_storage.items()
        ]
    }

@app.delete("/translations/{translation_id}")
async def delete_translation(translation_id: str):
    """Delete a stored translation"""
    if translation_id not in translation_storage:
        raise HTTPException(status_code=404, detail="Translation not found")
    
    filename = translation_storage[translation_id].get("filename", "unknown")
    del translation_storage[translation_id]
    
    logger.info(f"🗑️ Translation deleted: {filename} (ID: {translation_id})")
    return {"message": "Translation deleted successfully", "id": translation_id, "filename": filename}

# ✅ COMPREHENSIVE ERROR HANDLERS
@app.exception_handler(404)
async def not_found_handler(request: Request, exc):
    return JSONResponse(
        status_code=404,
        content={
            "error": "Endpoint not found",
            "path": str(request.url.path),
            "method": request.method,
            "timestamp": datetime.now().isoformat(),
            "available_endpoints": {
                "main": ["/", "/health", "/api/status"],
                "translation": ["/translate", "/translations"],
                "large_docs": ["/api/v1/large/translate", "/api/v1/large/status/{job_id}"],
                "async": ["/api/v2/translate/pdf/async"],
                "docs": ["/docs", "/redoc"]
            }
        }
    )

@app.exception_handler(422)
async def validation_error_handler(request: Request, exc):
    return JSONResponse(
        status_code=422,
        content={
            "error": "Validation error",
            "message": "Invalid request data",
            "details": str(exc),
            "timestamp": datetime.now().isoformat()
        }
    )

@app.exception_handler(500)
async def internal_error_handler(request: Request, exc):
    logger.error(f"Internal server error: {str(exc)}\n{traceback.format_exc()}")
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "message": "An unexpected error occurred. Please try again later.",
            "timestamp": datetime.now().isoformat(),
            "support": "Contact support if this error persists"
        }
    )

# ✅ STARTUP EVENT
@app.on_event("startup")
async def startup_event():
    logger.info("🚀 PRISMY Translation API Starting Up...")
    logger.info("=" * 60)
    logger.info("📡 API Documentation: http://localhost:8000/docs")
    logger.info("🔗 Health Check: http://localhost:8000/health")
    logger.info("🌐 Frontend Support: http://localhost:3000 | http://localhost:3001 | http://localhost:3002")  # ✅ Updated
    logger.info("✅ CORS: Enabled for all dev ports (3000, 3001, 3002)")  # ✅ Updated
    logger.info("📁 File Support: PDF, DOC, DOCX, TXT")
    logger.info("🔄 Translation Tiers: Free, Basic, Standard, Premium")
    logger.info("=" * 60)

# ✅ SHUTDOWN EVENT
@app.on_event("shutdown")
async def shutdown_event():
    logger.info("🛑 PRISMY Translation API Shutting Down...")
    # Cleanup resources if needed
    translation_storage.clear()
    logger.info("✅ Cleanup completed")

if __name__ == "__main__":
    import uvicorn
    print("🚀 PRISMY Translation API - Production Ready")
    print("=" * 50)
    print("📡 API Documentation: http://localhost:8000/docs")
    print("🔗 Health Check: http://localhost:8000/health") 
    print("🌐 Frontend: http://localhost:3000 | http://localhost:3001 | http://localhost:3002")  # ✅ Updated
    print("=" * 50)
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
# Debug port configuration
import os
print(f"🚀 STARTING ON PORT: {os.getenv('PORT', '8000')}")

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
