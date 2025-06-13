import os
import json
import logging
import asyncio
from typing import Dict, List, Optional, Tuple
from celery_app import app as celery_app
import redis
import PyPDF2
import pdfplumber
from PIL import Image
import pytesseract
import io
import re
from datetime import datetime

# Import từ local modules
from src.services.storage_service import StorageService
from src.services.queue.redis_client import get_redis_client
from src.core.models import TranslationJob, JobStatus, TranslationTier
from src.core.config import settings

# ============ THÊM MỚI: Import DOCX và FileExtractor ============
try:
    from docx import Document
    import zipfile
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX support disabled.")

# Configure Celery
celery_app.conf.update(
    task_serializer='json',
    accept_content=['json'],
    result_serializer='json', 
    timezone='UTC',
    enable_utc=True,
    task_track_started=True,
    task_time_limit=3600,
    task_soft_time_limit=3300,
)

logger = logging.getLogger(__name__)
storage_service = StorageService()
redis_client = get_redis_client()

# ============================================
# HELPER FUNCTIONS FOR REDIS HASH OPERATIONS
# ============================================

def get_job_data(job_id: str) -> Dict:
    """Get job data from Redis hash"""
    try:
        job_data = redis_client.hgetall(f"prismy:job:{job_id}")
        if not job_data:
            raise ValueError(f"Job {job_id} not found")
        
        # Parse JSON fields if they exist
        for key in ['extraction_result', 'translation_result']:
            if key in job_data and job_data[key]:
                try:
                    job_data[key] = json.loads(job_data[key])
                except json.JSONDecodeError:
                    pass
        
        return job_data
    except Exception as e:
        logger.error(f"Error getting job data for {job_id}: {e}")
        raise

def update_job_data(job_id: str, updates: Dict) -> None:
    """Update job data in Redis hash với Unicode handling"""
    try:
        # Prepare updates for Redis hash
        redis_updates = {}
        for key, value in updates.items():
            if isinstance(value, (dict, list)):
                redis_updates[key] = json.dumps(value, ensure_ascii=False)
            else:
                redis_updates[key] = str(value)
        
        # Add timestamp
        redis_updates['updated_at'] = datetime.utcnow().isoformat()
        
        # Update Redis hash
        redis_client.hset(f"prismy:job:{job_id}", mapping=redis_updates)
        
        # Set expiration (24 hours)
        redis_client.expire(f"prismy:job:{job_id}", 86400)
        
    except Exception as e:
        logger.error(f"Error updating job data for {job_id}: {e}")
        raise

def update_job_progress(job_id: str, status: str, progress: int, message: str = None) -> None:
    """Update job progress with status"""
    updates = {
        'status': status,
        'progress': progress
    }
    if message:
        updates['message'] = message
    
    update_job_data(job_id, updates)

def mark_job_failed(job_id: str, error: str) -> None:
    """Mark job as failed with error message"""
    update_job_data(job_id, {
        'status': JobStatus.FAILED.value,
        'error': error,
        'progress': 0
    })

# ============ CẬP NHẬT: File Type Detector ============
class FileTypeDetector:
    """Phát hiện chính xác loại file"""
    
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """Phát hiện loại file thực tế từ file path và extension"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"File không tồn tại: {file_path}")
                return 'unknown'
            
            # Lấy extension từ file path
            from pathlib import Path
            ext = Path(file_path).suffix.lower()
            
            # Ưu tiên extension trước
            if ext == '.docx':
                # Kiểm tra thêm bằng ZIP header
                try:
                    with open(file_path, 'rb') as f:
                        header = f.read(4)
                    if header.startswith(b'PK\x03\x04'):
                        logger.info(f"Detected DOCX from extension and header: {file_path}")
                        return 'docx'
                except:
                    pass
                logger.info(f"Detected DOCX from extension: {file_path}")
                return 'docx'
            
            elif ext == '.pdf':
                # Kiểm tra PDF header
                try:
                    with open(file_path, 'rb') as f:
                        header = f.read(8)
                    if header.startswith(b'%PDF-'):
                        logger.info(f"Detected PDF from extension and header: {file_path}")
                        return 'pdf'
                except:
                    pass
                logger.info(f"Detected PDF from extension: {file_path}")
                return 'pdf'
            
            elif ext in ['.txt', '.md', '.text']:
                logger.info(f"Detected TEXT from extension: {file_path}")
                return 'txt'
            
            # Nếu không có extension rõ ràng, đọc header
            try:
                with open(file_path, 'rb') as f:
                    header = f.read(8)
                
                # PDF check
                if header.startswith(b'%PDF-'):
                    logger.info(f"Detected PDF from header: {file_path}")
                    return 'pdf'
                
                # ZIP-based formats (DOCX)
                if header.startswith(b'PK\x03\x04'):
                    return FileTypeDetector._check_docx(file_path)
                
                # Text files
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        f.read(100)
                    logger.info(f"Detected TEXT from content: {file_path}")
                    return 'txt'
                except:
                    pass
                
            except Exception as e:
                logger.error(f"Error reading file header: {e}")
            
            logger.warning(f"Could not detect file type for: {file_path}")
            return 'unknown'
            
        except Exception as e:
            logger.error(f"File type detection error: {e}")
            return 'unknown'
    
    @staticmethod
    def _check_docx(file_path: str) -> str:
        """Kiểm tra file ZIP có phải DOCX không"""
        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                contents = zip_file.namelist()
                if 'word/document.xml' in contents:
                    logger.info(f"Confirmed DOCX structure: {file_path}")
                    return 'docx'
            logger.warning(f"ZIP file but not DOCX: {file_path}")
            return 'unknown'
        except Exception as e:
            logger.error(f"Error checking DOCX structure: {e}")
            return 'unknown'

# ============ THÊM MỚI: DOCX Processor ============
class DocxProcessor:
    """Xử lý file DOCX"""
    
    def __init__(self):
        self.name = "DocxProcessor"
    
    def process_docx(self, file_path: str) -> Dict:
        """Xử lý file DOCX và trả về dữ liệu có cấu trúc"""
        logger.info(f"Processing DOCX: {file_path}")
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            # Kiểm tra file DOCX hợp lệ
            if not self._is_valid_docx(file_path):
                raise ValueError(f"Invalid DOCX file: {file_path}")
            
            # Trích xuất bằng python-docx
            doc = Document(file_path)
            pages_data = []
            
            # Trích xuất paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    pages_data.append({
                        'page': 1,  # DOCX không có page concept rõ ràng
                        'text': text,
                        'method': 'docx_paragraph',
                        'element_type': 'paragraph',
                        'element_index': i
                    })
            
            # Trích xuất tables
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    pages_data.append({
                        'page': 1,
                        'text': table_text,
                        'method': 'docx_table',
                        'element_type': 'table',
                        'element_index': table_idx
                    })
            
            if not pages_data:
                raise ValueError("No content extracted from DOCX file")
            
            # Kết hợp tất cả text
            total_text = '\n\n'.join([p['text'] for p in pages_data])
            
            logger.info(f"DOCX processing completed: {len(pages_data)} elements, {len(total_text)} characters")
            
            return {
                'pages': pages_data,
                'total_pages': 1,  # DOCX coi như 1 trang
                'total_characters': len(total_text),
                'extraction_methods': ['docx'],
                'document_type': 'docx'
            }
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise
    
    def _is_valid_docx(self, file_path: str) -> bool:
        """Kiểm tra DOCX hợp lệ"""
        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                required_files = ['word/document.xml', '[Content_Types].xml']
                zip_contents = zip_file.namelist()
                return all(req in zip_contents for req in required_files)
        except Exception as e:
            logger.error(f"DOCX validation error: {e}")
            return False
    
    def _extract_table_text(self, table) -> str:
        """Trích xuất text từ table"""
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append('\t'.join(row_text))
        return '\n'.join(table_text)

# ============ CẬP NHẬT: PDF Processor ============
class PDFProcessor:
    """Advanced PDF processing with text extraction and OCR support"""
    
    def __init__(self):
        self.min_text_length = 10
        self.ocr_languages = ['eng', 'vie']
    
    def extract_text_pypdf2(self, file_path: str) -> List[Dict]:
        """Extract text using PyPDF2"""
        pages_data = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text and len(text.strip()) > self.min_text_length:
                        pages_data.append({
                            'page': page_num + 1,
                            'text': text.strip(),
                            'method': 'pypdf2',
                            'element_type': 'page'
                        })
                        
        except Exception as e:
            logger.error(f"PyPDF2 extraction error: {str(e)}")
            
        return pages_data
    
    def extract_text_pdfplumber(self, file_path: str) -> List[Dict]:
        """Extract text using pdfplumber (better for tables)"""
        pages_data = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text()
                    
                    # Extract tables
                    tables = page.extract_tables()
                    table_text = ""
                    
                    if tables:
                        for table in tables:
                            for row in table:
                                if row:
                                    table_text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                    
                    combined_text = f"{text or ''}\n{table_text}".strip()
                    
                    if combined_text and len(combined_text) > self.min_text_length:
                        pages_data.append({
                            'page': page_num + 1,
                            'text': combined_text,
                            'method': 'pdfplumber',
                            'element_type': 'page',
                            'has_tables': bool(tables)
                        })
                        
        except Exception as e:
            logger.error(f"pdfplumber extraction error: {str(e)}")
            
        return pages_data
    
    def process_pdf(self, file_path: str) -> Dict:
        """Main PDF processing method with fallback strategies"""
        logger.info(f"Processing PDF: {file_path}")
        
        # Try pdfplumber first (better for most cases)
        pages_data = self.extract_text_pdfplumber(file_path)
        
        # If not enough content, try PyPDF2
        if not pages_data or sum(len(p['text']) for p in pages_data) < 100:
            logger.info("Trying PyPDF2 extraction...")
            pypdf2_pages = self.extract_text_pypdf2(file_path)
            pages_data.extend(pypdf2_pages)
        
        # Merge and clean results
        final_pages = self._merge_pages_data(pages_data)
        
        if not final_pages:
            raise ValueError("No content extracted from PDF file")
        
        logger.info(f"PDF processing completed: {len(final_pages)} pages")
        
        return {
            'pages': final_pages,
            'total_pages': len(final_pages),
            'total_characters': sum(len(p['text']) for p in final_pages),
            'extraction_methods': list(set(p['method'] for p in final_pages)),
            'document_type': 'pdf'
        }
    
    def _merge_pages_data(self, pages_data: List[Dict]) -> List[Dict]:
        """Merge and deduplicate pages data"""
        merged = {}
        
        for page in pages_data:
            page_num = page['page']
            if page_num not in merged or len(page['text']) > len(merged[page_num]['text']):
                merged[page_num] = page
        
        return [merged[k] for k in sorted(merged.keys())]

# ============ THÊM MỚI: Text Processor ============
class TextProcessor:
    """Xử lý file text thuần"""
    
    def process_text(self, file_path: str) -> Dict:
        """Xử lý file text"""
        logger.info(f"Processing TEXT: {file_path}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.info(f"Successfully read text file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Could not decode text file with any encoding")
            
            if not content.strip():
                raise ValueError("Text file is empty")
            
            # Chia thành paragraphs
            paragraphs = content.split('\n\n')
            pages_data = []
            
            for i, para in enumerate(paragraphs):
                if para.strip():
                    pages_data.append({
                        'page': 1,
                        'text': para.strip(),
                        'method': 'text_paragraph',
                        'element_type': 'paragraph',
                        'paragraph_index': i
                    })
            
            if not pages_data:
                # If no paragraphs, treat entire content as one element
                pages_data.append({
                    'page': 1,
                    'text': content.strip(),
                    'method': 'text_content',
                    'element_type': 'content'
                })
            
            logger.info(f"Text processing completed: {len(pages_data)} paragraphs")
            
            return {
                'pages': pages_data,
                'total_pages': 1,
                'total_characters': len(content),
                'extraction_methods': ['text'],
                'document_type': 'txt'
            }
            
        except Exception as e:
            logger.error(f"Text processing error: {e}")
            raise

# Initialize processors
pdf_processor = PDFProcessor()
docx_processor = DocxProcessor()
text_processor = TextProcessor()

# ============ CẬP NHẬT: Extract Text Task ============
@celery_app.task(name='prismy_tasks.extract_text')
def extract_text(file_path: str, file_type: str = None) -> List[Dict]:
    """Extract text from uploaded file with auto-detection - Returns chunks directly"""
    try:
        logger.info(f"Starting extraction for {file_path}")
        
        # Auto-detect file type if not provided or incorrect
        detected_type = FileTypeDetector.detect_file_type(file_path)
        
        # Use detected type if provided type is wrong or missing
        if not file_type or file_type == 'unknown' or file_type != detected_type:
            file_type = detected_type
            logger.info(f"Using detected file type: {file_type}")
        
        extracted_data = {}
        
        # Process based on file type
        if file_type == 'pdf':
            extracted_data = pdf_processor.process_pdf(file_path)
        elif file_type == 'docx':
            if not DOCX_AVAILABLE:
                raise ImportError("DOCX support not available. Install: pip install python-docx")
            extracted_data = docx_processor.process_docx(file_path)
        elif file_type in ['txt', 'text']:
            extracted_data = text_processor.process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Format for translation - convert pages to chunks
        text_chunks = []
        for page in extracted_data['pages']:
            chunks = _split_text_into_chunks(page['text'], max_chunk_size=1000)
            for chunk_idx, chunk in enumerate(chunks):
                text_chunks.append({
                    'page': page['page'],
                    'text': chunk,
                    'method': page['method'],
                    'chunk_index': chunk_idx,
                    'element_type': page.get('element_type', 'text')
                })
        
        logger.info(f"Extraction completed: {len(text_chunks)} chunks from {file_type}")
        return text_chunks
        
    except Exception as e:
        logger.error(f"Extraction failed for {file_path}: {str(e)}")
        raise

def _split_text_into_chunks(text: str, max_chunk_size: int = 1000) -> List[str]:
    """Split text into manageable chunks for translation"""
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 < max_chunk_size:  # +2 for \n\n
            if current_chunk:
                current_chunk += "\n\n" + para
            else:
                current_chunk = para
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            # If paragraph is too long, split by sentences
            if len(para) > max_chunk_size:
                sentences = re.split(r'(?<=[.!?])\s+', para)
                sentence_chunk = ""
                
                for sentence in sentences:
                    if len(sentence_chunk) + len(sentence) + 1 < max_chunk_size:
                        if sentence_chunk:
                            sentence_chunk += " " + sentence
                        else:
                            sentence_chunk = sentence
                    else:
                        if sentence_chunk:
                            chunks.append(sentence_chunk.strip())
                        sentence_chunk = sentence
                
                if sentence_chunk:
                    current_chunk = sentence_chunk
                else:
                    current_chunk = ""
            else:
                current_chunk = para
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

# ============ CẬP NHẬT: Translate Chunks Task ============
@celery_app.task(name='prismy_tasks.translate_chunks')
def translate_chunks(extracted_chunks: List[Dict], target_language: str, tier: str, job_id: str) -> List[Dict]:
    """Translate text chunks using translation APIs"""
    try:
        logger.info(f"Starting translation for job {job_id}: {len(extracted_chunks)} chunks to {target_language}")
        
        # Import translation manager
        from src.services.translation_manager import TranslationManager
        translation_manager = TranslationManager()
        provider = translation_manager.get_provider(tier)
        
        translated_chunks = []
        total_chunks = len(extracted_chunks)
        
        # Translate each chunk
        for i, chunk in enumerate(extracted_chunks):
            try:
                # Translate the chunk
                translated_text = asyncio.run(
                    provider.translate(chunk['text'], target_language, 'auto')
                )
                
                translated_chunk = {
                    'page': chunk.get('page', 1),
                    'original': chunk['text'],
                    'translated': translated_text,
                    'method': chunk.get('method', 'unknown'),
                    'element_type': chunk.get('element_type', 'text')
                }
                translated_chunks.append(translated_chunk)
                
                logger.info(f"Translated chunk {i+1}/{total_chunks}")
                
            except Exception as e:
                logger.error(f"Translation failed for chunk {i+1}: {str(e)}")
                # Add original text as fallback
                translated_chunk = {
                    'page': chunk.get('page', 1),
                    'original': chunk['text'],
                    'translated': f"[Translation Error: {str(e)}] {chunk['text']}",
                    'method': chunk.get('method', 'unknown'),
                    'element_type': chunk.get('element_type', 'text')
                }
                translated_chunks.append(translated_chunk)
        
        logger.info(f"Translation completed for job {job_id}: {len(translated_chunks)} chunks")
        return translated_chunks
        
    except Exception as e:
        logger.error(f"Translation failed for job {job_id}: {str(e)}")
        raise

# ============ CẬP NHẬT: Reconstruct Document Task ============
@celery_app.task(name='prismy_tasks.reconstruct_document')
def reconstruct_document(translated_chunks: List[Dict], job_id: str, output_format: str) -> str:
    """Reconstruct translated document với Unicode fix"""
    try:
        logger.info(f"Starting reconstruction for job {job_id}")
        
        if not translated_chunks:
            logger.warning("No chunks to reconstruct!")
            return _create_empty_file(job_id)
        
        # Group chunks by page
        pages = {}
        for chunk in translated_chunks:
            page_num = chunk.get('page', 1)
            if page_num not in pages:
                pages[page_num] = []
            
            # Get translated text với Unicode decode
            translated_text = chunk.get('translated', chunk.get('original', chunk.get('text', '')))
            
            # Fix Unicode escape sequences
            if translated_text and isinstance(translated_text, str):
                try:
                    # Try to decode Unicode escape sequences
                    if '\\u' in translated_text:
                        import codecs
                        decoded_text = codecs.decode(translated_text, 'unicode_escape')
                        translated_text = decoded_text
                except Exception as e:
                    logger.warning(f"Unicode decode failed for chunk: {e}")
                
                pages[page_num].append(translated_text)
        
        # Create output document
        output_content = "TRANSLATED DOCUMENT\n"
        output_content += "=" * 50 + "\n\n"
        
        for page_num in sorted(pages.keys()):
            if len(pages) > 1:
                output_content += f"PAGE {page_num}\n"
                output_content += "-" * 30 + "\n"
            
            page_content = "\n\n".join(pages[page_num])
            output_content += page_content + "\n\n"
        
        # Save output file với UTF-8 encoding
        output_filename = f"translated_{job_id}.txt"
        output_path = os.path.join(settings.OUTPUT_DIR, output_filename)
        
        os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output_content)
        
        logger.info(f"Reconstruction completed for job {job_id}: {output_filename}")
        return output_filename
        
    except Exception as e:
        logger.error(f"Reconstruction failed for job {job_id}: {str(e)}")
        raise

def _create_empty_file(job_id: str) -> str:
    """Create empty file when no content"""
    output_filename = f"translated_{job_id}.txt"
    output_path = os.path.join(settings.OUTPUT_DIR, output_filename)
    
    os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("ERROR: No content was translated. Please check the source file.\n")
    
    return output_filename

# ============ FIXED: Process Translation Task - NO .get() CALLS ============
@celery_app.task(name='prismy_tasks.process_translation_sync')
def process_translation_sync(job_id: str, file_path: str, file_type: str, target_lang: str, tier: str) -> str:
    """
    Synchronous translation process - Executes all steps directly without .get() calls
    This fixes the Celery "Never call result.get() within a task" error
    """
    try:
        logger.info(f"🚀 Starting synchronous translation for job {job_id}")
        logger.info(f"📁 File: {file_path}")
        logger.info(f"🎯 File type: {file_type} → Translation: {target_lang}, Tier: {tier}")
        
        update_job_progress(job_id, JobStatus.PROCESSING.value, 10, "Starting translation...")
        
        # Step 1: Extract text directly
        logger.info(f"📖 Step 1: Extracting text from file...")
        try:
            extracted_chunks = extract_text(file_path, file_type)
            
            if not extracted_chunks:
                raise ValueError("No content extracted from file")
            
            logger.info(f"✅ Extraction completed: {len(extracted_chunks)} chunks")
            update_job_progress(job_id, JobStatus.PROCESSING.value, 40, f"Extracted {len(extracted_chunks)} chunks")
            
        except Exception as e:
            logger.error(f"❌ Extraction failed: {str(e)}")
            raise
        
        # Step 2: Translate chunks directly
        logger.info(f"🔄 Step 2: Translating {len(extracted_chunks)} chunks...")
        try:
            translated_chunks = translate_chunks(extracted_chunks, target_lang, tier, job_id)
            
            if not translated_chunks:
                raise ValueError("Translation failed - no results")
            
            logger.info(f"✅ Translation completed: {len(translated_chunks)} chunks")
            update_job_progress(job_id, JobStatus.PROCESSING.value, 80, "Translation completed")
            
        except Exception as e:
            logger.error(f"❌ Translation failed: {str(e)}")
            raise
        
        # Step 3: Reconstruct document directly
        logger.info(f"📄 Step 3: Reconstructing document...")
        try:
            output_filename = reconstruct_document(translated_chunks, job_id, 'txt')
            
            if not output_filename:
                raise ValueError("Document reconstruction failed")
            
            # Final status update
            update_job_data(job_id, {
                'status': JobStatus.COMPLETED.value,
                'progress': 100,
                'output_file': output_filename,
                'output_path': os.path.join(settings.OUTPUT_DIR, output_filename),
                'message': 'Translation completed successfully!'
            })
            
            logger.info(f"🎉 Translation process completed for job {job_id}")
            logger.info(f"📥 Output file: {output_filename}")
            
            return output_filename
            
        except Exception as e:
            logger.error(f"❌ Reconstruction failed: {str(e)}")
            raise
        
    except Exception as e:
        error_msg = f"Translation failed: {str(e)}"
        logger.error(f"❌ Job {job_id} failed: {error_msg}")
        mark_job_failed(job_id, error_msg)
        raise

# ============ LEGACY: Keep old task for compatibility ============
@celery_app.task(name='prismy_tasks.process_translation')
def process_translation(job_id: str, file_path: str, file_type: str, target_lang: str, tier: str) -> str:
    """
    Legacy task - redirects to sync version
    """
    logger.info(f"Legacy process_translation called, redirecting to sync version")
    return process_translation_sync(job_id, file_path, file_type, target_lang, tier)