# src/processors/docx_extractor.py
import logging
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
import zipfile
import os

logger = logging.getLogger(__name__)

class DocxExtractor:
    """Trích xuất văn bản từ file DOCX"""
    
    def __init__(self):
        self.name = "DocxExtractor"
    
    def extract(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Trích xuất văn bản từ file DOCX
        
        Args:
            file_path: Đường dẫn đến file DOCX
            
        Returns:
            Danh sách các đoạn văn bản với thông tin metadata
        """
        try:
            # Kiểm tra file có tồn tại không
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
            
            # Kiểm tra file có phải DOCX hợp lệ không
            if not self._is_valid_docx(file_path):
                raise ValueError(f"File DOCX không hợp lệ: {file_path}")
            
            # Trích xuất văn bản bằng python-docx
            doc = Document(file_path)
            chunks = []
            
            logger.info(f"Bắt đầu trích xuất DOCX: {file_path}")
            
            # Trích xuất các đoạn văn
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:  # Chỉ thêm đoạn văn không rỗng
                    chunks.append({
                        'chunk_id': f"para_{i}",
                        'text': text,
                        'type': 'paragraph',
                        'metadata': {
                            'paragraph_index': i,
                            'style': paragraph.style.name if paragraph.style else 'Normal'
                        }
                    })
            
            # Trích xuất bảng biểu
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    chunks.append({
                        'chunk_id': f"table_{table_idx}",
                        'text': table_text,
                        'type': 'table',
                        'metadata': {
                            'table_index': table_idx,
                            'rows': len(table.rows),
                            'cols': len(table.columns) if table.rows else 0
                        }
                    })
            
            logger.info(f"Trích xuất thành công {len(chunks)} đoạn từ DOCX: {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Lỗi trích xuất DOCX: {str(e)}")
            raise
    
    def _is_valid_docx(self, file_path: str) -> bool:
        """Kiểm tra xem file có phải DOCX hợp lệ không"""
        try:
            # Kiểm tra phần mở rộng file
            if not file_path.lower().endswith('.docx'):
                return False
            
            # Kiểm tra xem có phải file ZIP hợp lệ không (DOCX là container ZIP)
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Kiểm tra các file bắt buộc của DOCX
                required_files = ['word/document.xml', '[Content_Types].xml']
                zip_contents = zip_file.namelist()
                
                for required_file in required_files:
                    if required_file not in zip_contents:
                        return False
                
                return True
                
        except (zipfile.BadZipFile, FileNotFoundError):
            return False
    
    def _extract_table_text(self, table) -> str:
        """Trích xuất văn bản từ bảng và định dạng"""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            
            # Nối các ô bằng tab
            table_text.append('\t'.join(row_text))
        
        # Nối các hàng bằng xuống dòng
        return '\n'.join(table_text)

# Hàm tạo extractor
def create_docx_extractor():
    """Tạo bộ trích xuất DOCX"""
    return DocxExtractor()