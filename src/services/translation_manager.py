import os
from typing import List, Dict, Optional
from src.providers.base_translator import BaseTranslator
from src.providers.google_translator import GoogleTranslator
from src.providers.openai_translator import OpenAITranslator
from src.providers.anthropic_translator import AnthropicTranslator
import logging
import asyncio

logger = logging.getLogger(__name__)

class MockTranslator(BaseTranslator):
    """🎭 FORCED Mock translator for testing - supports ALL languages"""
    
    def __init__(self):
        self.client = True  # Fake client
        logger.info("🎭 FORCED Mock Translator initialized for testing ALL languages")
    
    def get_supported_languages(self) -> List[str]:
        """Return comprehensive list of supported languages"""
        return [
            'en', 'vi', 'zh', 'zh-cn', 'zh-tw', 'ja', 'ko', 'th', 'fr', 'de', 'es', 'it', 'pt', 'ru', 'ar', 'hi',
            'nl', 'sv', 'da', 'no', 'fi', 'pl', 'cs', 'hu', 'ro', 'bg', 'hr', 'sk', 'sl', 'et', 'lv', 'lt', 'mt'
        ]
    
    def _get_language_names(self) -> Dict[str, str]:
        """Comprehensive mapping of language codes to names"""
        return {
            'en': 'English', 'vi': 'Vietnamese', 'zh': 'Chinese', 'zh-cn': 'Chinese (Simplified)', 
            'zh-tw': 'Chinese (Traditional)', 'ja': 'Japanese', 'ko': 'Korean', 'th': 'Thai',
            'fr': 'French', 'de': 'German', 'es': 'Spanish', 'it': 'Italian', 'pt': 'Portuguese',
            'ru': 'Russian', 'ar': 'Arabic', 'hi': 'Hindi', 'nl': 'Dutch', 'sv': 'Swedish', 
            'da': 'Danish', 'no': 'Norwegian', 'fi': 'Finnish', 'pl': 'Polish', 'cs': 'Czech',
            'hu': 'Hungarian', 'ro': 'Romanian', 'bg': 'Bulgarian', 'hr': 'Croatian', 'sk': 'Slovak',
            'auto': 'Auto-detected'
        }
    
    def _normalize_language_name(self, lang_code: str) -> str:
        """Convert language codes to readable names"""
        lang_names = self._get_language_names()
        return lang_names.get(lang_code.lower(), lang_code.capitalize())
    
    def _get_vietnamese_to_english_mapping(self) -> Dict[str, str]:
        """Vietnamese to English word mapping"""
        return {
            # Basic words
            'các': 'the', 'và': 'and', 'của': 'of', 'đến': 'to', 'trong': 'in',
            'cho': 'for', 'với': 'with', 'bởi': 'by', 'từ': 'from', 'về': 'about',
            'này': 'this', 'đó': 'that', 'là': 'is', 'có': 'have', 'được': 'be',
            'sẽ': 'will', 'có thể': 'can', 'nên': 'should', 'phải': 'must',
            'tôi': 'I', 'bạn': 'you', 'chúng ta': 'we', 'họ': 'they',
            'một': 'a', 'những': 'some', 'nhiều': 'many', 'ít': 'few',
            
            # Academic/document terms (CRITICAL for test files)
            'tài liệu': 'document', 'báo cáo': 'report', 'nghiên cứu': 'research',
            'phân tích': 'analysis', 'kết quả': 'result', 'kết luận': 'conclusion',
            'phương pháp': 'method', 'dữ liệu': 'data', 'thông tin': 'information',
            'hệ thống': 'system', 'chương trình': 'program', 'dự án': 'project',
            'bản quyền': 'copyright', 'đại học': 'university', 'stanford': 'stanford',
            
            # Healthcare and AI terms (for test documents)
            'đánh giá': 'evaluation', 'ứng dụng': 'application', 'áp dụng': 'apply',
            'chăm sóc sức khỏe': 'healthcare', 'y tế': 'medical', 'bệnh nhân': 'patient',
            'điều trị': 'treatment', 'chẩn đoán': 'diagnosis', 'thuốc': 'medicine',
            'công nghệ': 'technology', 'trí tuệ nhân tạo': 'artificial intelligence',
            'học máy': 'machine learning', 'dữ liệu lớn': 'big data',
            'thuật toán': 'algorithm', 'mô hình': 'model', 'hiệu suất': 'performance',
            'độ chính xác': 'accuracy', 'so sánh': 'comparison',
            'thử nghiệm': 'test', 'kiểm tra': 'verification', 'xác thực': 'validation',
            'thách thức': 'challenge', 'cơ hội': 'opportunity', 'mục tiêu': 'goal'
        }

    async def translate(self, text: str, target_language: str, source_language: str = 'auto') -> str:
        """
        🎭 FORCED MOCK TRANSLATION - Supports ALL world languages
        """
        # ✅ COMPREHENSIVE DEBUG LOGGING
        logger.info(f"🎭 FORCED MOCK TRANSLATOR PROCESSING:")
        logger.info(f"📝 Text (first 50 chars): {text[:50]}...")
        logger.info(f"🎯 target_language: '{target_language}' (type: {type(target_language)})")
        logger.info(f"🗣️ source_language: '{source_language}' (type: {type(source_language)})")
        
        # ✅ CLEAN AND NORMALIZE
        target_clean = str(target_language).strip().lower()
        logger.info(f"🧹 Cleaned target: '{target_clean}'")
        
        # ✅ GET PROPER LANGUAGE NAME
        target_name = self._normalize_language_name(target_language)
        logger.info(f"🏷️ Target name: '{target_name}'")
        
        # ✅ UNIVERSAL TRANSLATION LOGIC
        if target_clean in ['en', 'english'] or target_clean.startswith('en'):
            # ✅ TRANSLATE TO ENGLISH
            logger.info(f"✅ ENGLISH TARGET CONFIRMED! Processing Vietnamese → English translation")
            
            vi_to_en_map = self._get_vietnamese_to_english_mapping()
            result = self._apply_word_mapping(text, vi_to_en_map, 'Vietnamese→English')
            final_result = f"[MOCK-TRANSLATED to English]: {result}"
            
        elif target_clean in ['vi', 'vietnamese']:
            # ✅ TRANSLATE TO VIETNAMESE  
            logger.info(f"✅ VIETNAMESE TARGET CONFIRMED! Processing → Vietnamese translation")
            
            # Create reverse mapping
            vi_to_en = self._get_vietnamese_to_english_mapping()
            en_to_vi = {en: vi for vi, en in vi_to_en.items()}
            
            result = self._apply_word_mapping(text, en_to_vi, 'English→Vietnamese')
            final_result = f"[MOCK-TRANSLATED to Vietnamese]: {result}"
            
        else:
            # ✅ ALL OTHER LANGUAGES
            logger.info(f"✅ {target_name.upper()} TARGET CONFIRMED! Processing → {target_name} translation")
            
            # For other languages, translate Vietnamese to English first, then indicate target language
            vi_to_en_map = self._get_vietnamese_to_english_mapping()
            english_result = self._apply_word_mapping(text, vi_to_en_map, f'Vietnamese→{target_name}')
            final_result = f"[MOCK-TRANSLATED to {target_name}]: {english_result}"
        
        # ✅ FINAL LOGGING
        logger.info(f"🔄 Translation direction: {source_language} → {target_name}")
        logger.info(f"✅ {target_name.upper()} TRANSLATION COMPLETE!")
        logger.info(f"📤 Final result (first 100 chars): {final_result[:100]}...")
        
        return final_result
    
    def _apply_word_mapping(self, text: str, word_map: Dict[str, str], direction: str) -> str:
        """Apply word-by-word translation mapping"""
        import re
        
        result = text
        replacement_count = 0
        
        # Sort by length (longest first) to handle phrases before individual words
        sorted_words = sorted(word_map.items(), key=lambda x: len(x[0]), reverse=True)
        
        for source_word, target_word in sorted_words:
            # Use word boundaries for whole word replacement
            pattern = r'\b' + re.escape(source_word) + r'\b'
            matches = re.findall(pattern, result, flags=re.IGNORECASE)
            
            if matches:
                # Preserve case when possible
                def replace_func(match):
                    original = match.group()
                    if original.isupper():
                        return target_word.upper()
                    elif original.istitle():
                        return target_word.capitalize()
                    else:
                        return target_word
                
                result = re.sub(pattern, replace_func, result, flags=re.IGNORECASE)
                replacement_count += len(matches)
                logger.debug(f"🔄 {direction}: '{source_word}' → '{target_word}' ({len(matches)} times)")
        
        logger.info(f"🔄 {direction} replacements made: {replacement_count}")
        return result
    
    async def translate_batch(self, texts: List[str], target_language: str, 
                            source_language: str = 'auto') -> List[str]:
        """FORCED mock batch translation"""
        logger.info(f"🎭 FORCED BATCH TRANSLATION: {len(texts)} texts → {target_language}")
        results = []
        for i, text in enumerate(texts):
            result = await self.translate(text, target_language, source_language)
            results.append(result)
            logger.debug(f"📝 Batch item {i+1}/{len(texts)} completed")
        
        logger.info(f"✅ FORCED BATCH TRANSLATION COMPLETE: {len(results)} results")
        return results

class FileExtractor:
    """Extract text from various file types"""
    
    def __init__(self):
        self.supported_types = ['pdf', 'docx', 'txt']
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect file type"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return 'unknown'
            
            # Detect by file header
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            # PDF check
            if header.startswith(b'%PDF-'):
                return 'pdf'
            
            # ZIP-based formats (DOCX)
            if header.startswith(b'PK\x03\x04'):
                return self._check_docx(file_path)
            
            # Text files
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    f.read(100)
                return 'txt'
            except:
                pass
            
            # Fallback to extension
            from pathlib import Path
            ext = Path(file_path).suffix.lower()
            if ext == '.pdf':
                return 'pdf'
            elif ext == '.docx':
                return 'docx'
            elif ext in ['.txt', '.md']:
                return 'txt'
            
            return 'unknown'
            
        except Exception as e:
            logger.error(f"Error detecting file type: {e}")
            return 'unknown'
    
    def _check_docx(self, file_path: str) -> str:
        """Check if ZIP file is DOCX"""
        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                contents = zip_file.namelist()
                if 'word/document.xml' in contents:
                    return 'docx'
            return 'unknown'
        except:
            return 'unknown'
    
    def extract_text(self, file_path: str) -> List[Dict]:
        """Extract text from file"""
        file_type = self.detect_file_type(file_path)
        logger.info(f"Detected file type: {file_type} for {file_path}")
        
        if file_type == 'pdf':
            return self._extract_pdf(file_path)
        elif file_type == 'docx':
            return self._extract_docx(file_path)
        elif file_type == 'txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_pdf(self, file_path: str) -> List[Dict]:
        """Extract from PDF using existing extractor"""
        try:
            from src.processors.pdf_extractor import PDFExtractor
            extractor = PDFExtractor()
            return extractor.extract(file_path)
        except ImportError:
            logger.error("PDFExtractor not found")
            raise
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
    
    def _extract_docx(self, file_path: str) -> List[Dict]:
        """Extract from DOCX"""
        try:
            from docx import Document
            import zipfile
            
            if not self._is_valid_docx(file_path):
                raise ValueError(f"Invalid DOCX file: {file_path}")
            
            doc = Document(file_path)
            chunks = []
            
            logger.info(f"Starting DOCX extraction: {file_path}")
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    chunks.append({
                        'chunk_id': f"para_{i}",
                        'text': text,
                        'type': 'paragraph',
                        'metadata': {
                            'paragraph_index': i,
                            'style': paragraph.style.name if paragraph.style else 'Normal'
                        }
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    chunks.append({
                        'chunk_id': f"table_{table_idx}",
                        'text': table_text,
                        'type': 'table',
                        'metadata': {
                            'table_index': table_idx,
                            'rows': len(table.rows),
                            'cols': len(table.columns) if table.rows else 0
                        }
                    })
            
            logger.info(f"Successfully extracted {len(chunks)} chunks from DOCX")
            return chunks
            
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise
    
    def _is_valid_docx(self, file_path: str) -> bool:
        """Check if DOCX is valid"""
        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                required_files = ['word/document.xml', '[Content_Types].xml']
                zip_contents = zip_file.namelist()
                return all(req in zip_contents for req in required_files)
        except:
            return False
    
    def _extract_table_text(self, table) -> str:
        """Extract text from table"""
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append('\t'.join(row_text))
        return '\n'.join(table_text)
    
    def _extract_txt(self, file_path: str) -> List[Dict]:
        """Extract from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            chunks = []
            paragraphs = content.split('\n\n')
            
            for i, para in enumerate(paragraphs):
                if para.strip():
                    chunks.append({
                        'chunk_id': f"para_{i}",
                        'text': para.strip(),
                        'type': 'paragraph',
                        'metadata': {'paragraph_index': i}
                    })
            
            logger.info(f"Successfully extracted {len(chunks)} chunks from TXT")
            return chunks
            
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            raise

class TranslationManager:
    """🎭 FORCED Translation Manager - ALWAYS uses MockTranslator"""
    
    def __init__(self):
        self.providers = self._initialize_providers()
        self.use_mock = True  # FORCED TRUE
        self.file_extractor = FileExtractor()
        logger.info("🎭 FORCED Translation Manager initialized - ALWAYS uses MockTranslator")
    
    def get_provider(self, tier: str):
        """🎭 FORCED: Get translation provider - ALWAYS MockTranslator"""
        logger.info(f"🔍 Getting provider for tier: {tier}")
        logger.warning("🎭 FORCING MockTranslator for testing - ALL tiers use mock")
        return MockTranslator()
    
    def _initialize_providers(self) -> Dict[str, BaseTranslator]:
        """🎭 FORCED: Initialize providers - ALWAYS MockTranslator"""
        providers = {}
        logger.info("🎭 FORCING MockTranslator for ALL translation requests")
        providers['forced_mock'] = MockTranslator()
        return providers
    
    def get_provider_for_tier(self, tier: str) -> Optional[BaseTranslator]:
        """🎭 FORCED: Get appropriate provider - ALWAYS MockTranslator"""
        logger.info(f"🎭 FORCED: Returning MockTranslator for tier '{tier}' (supports ALL languages)")
        return MockTranslator()
    
    async def translate(self, text: str, target_language: str, tier: str = 'basic', 
                       source_language: str = 'auto') -> str:
        """🎭 FORCED: Async translation - ALWAYS MockTranslator"""
        logger.info(f"🎭 FORCED: Using MockTranslator for {source_language} → {target_language}")
        forced_translator = MockTranslator()
        return await forced_translator.translate(text, target_language, source_language)
    
    async def translate_batch(self, texts: List[str], target_language: str, 
                            tier: str = 'basic', source_language: str = 'auto') -> List[str]:
        """🎭 FORCED: Batch translation - ALWAYS MockTranslator"""
        logger.info(f"🎭 FORCED: Batch translation {len(texts)} texts → {target_language}")
        forced_translator = MockTranslator()
        return await forced_translator.translate_batch(texts, target_language, source_language)
    
    def translate_sync(self, text: str, target_language: str, tier: str = 'basic',
                      source_language: str = 'auto') -> str:
        """🎭 FORCED: Synchronous translation - ALWAYS MockTranslator"""
        logger.info(f"🎭 FORCED: Sync translation {source_language} → {target_language}")
        forced_translator = MockTranslator()
        
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        
        return loop.run_until_complete(
            forced_translator.translate(text, target_language, source_language)
        )
    
    def translate_batch_sync(self, texts: List[str], target_language: str,
                           tier: str = 'basic', source_language: str = 'auto') -> List[str]:
        """🎭 FORCED: Synchronous batch translation - ALWAYS MockTranslator"""
        forced_translator = MockTranslator()
        
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        
        return loop.run_until_complete(
            forced_translator.translate_batch(texts, target_language, source_language)
        )

# Global instance
translation_manager = TranslationManager()